import os
import difflib
from typing import List, Tuple
from docx import Document
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from docdiff.core.aligner import ParagraphAligner

def _extract_font_props(p: Paragraph) -> dict:

    font_name = None
    font_size = None
    bold = None
    italic = None

    for run in p.runs:
        if run.font.name:
            font_name = run.font.name
        if not font_name and run._element.rPr is not None and run._element.rPr.rFonts is not None:
            rFonts = run._element.rPr.rFonts
            font_name = rFonts.get(qn('w:eastAsia')) or rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi'))
        if run.font.size:
            font_size = run.font.size
        if run.bold is not None:
            bold = run.bold
        if run.italic is not None:
            italic = run.italic
        if font_name or font_size:
            break

    return {
        'name': font_name,
        'size': font_size,
        'bold': bold,
        'italic': italic
    }

def _apply_font_props(dst_run, font_props: dict):
    if not font_props:
        return
    font_name = font_props.get('name')
    font_size = font_props.get('size')
    bold = font_props.get('bold')
    italic = font_props.get('italic')

    if font_name:
        dst_run.font.name = font_name
        rPr = dst_run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)

    if font_size:
        dst_run.font.size = font_size
    if bold is not None and not dst_run.bold:
        dst_run.bold = bold
    if italic is not None and not dst_run.italic:
        dst_run.italic = italic

def _clear_paragraph_runs(p: Paragraph):
    """Removes all run elements from paragraph while preserving paragraph properties (<w:pPr>)."""
    for child in list(p._element):
        if child.tag.endswith('r'):
            p._element.remove(child)

class DocDiffEngine:
    """
    Core document comparison engine using python-docx with in-place paragraph mutation.
    Preserves 100% of cover pages, images, headers, footers, and section formatting.
    """

    def __init__(self, old_path: str, new_path: str, output_path: str, show_deletions: bool = True):
        self.old_path = old_path
        self.new_path = new_path
        self.output_path = output_path
        self.show_deletions = show_deletions
        
        self.red_color = RGBColor(255, 0, 0)

    def _compare_text(self, old_text: str, new_text: str) -> List[Tuple[str, str]]:
        matcher = difflib.SequenceMatcher(None, old_text, new_text)
        results = []
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                results.append(('equal', new_text[j1:j2]))
            elif tag == 'insert':
                results.append(('insert', new_text[j1:j2]))
            elif tag == 'delete':
                if self.show_deletions:
                    results.append(('delete', old_text[i1:i2]))
            elif tag == 'replace':
                if self.show_deletions:
                    results.append(('delete', old_text[i1:i2]))
                results.append(('insert', new_text[j1:j2]))
        return results

    def execute(self) -> str:
        if not os.path.exists(self.old_path):
            raise FileNotFoundError(f"原文件不存在: {self.old_path}")
        if not os.path.exists(self.new_path):
            raise FileNotFoundError(f"新文件不存在: {self.new_path}")

        old_doc = Document(self.old_path)
        out_doc = Document(self.new_path) # Modify out_doc IN-PLACE!

        old_paragraphs = old_doc.paragraphs
        new_paragraphs = out_doc.paragraphs

        old_texts = [p.text for p in old_paragraphs]
        new_texts = [p.text for p in new_paragraphs]

        aligned_pairs = ParagraphAligner.align_paragraphs(old_texts, new_texts)

        last_seen_new_p = None

        for tag, old_idx, new_idx in aligned_pairs:
            if tag in ('equal', 'replace'):
                old_p = old_paragraphs[old_idx]
                new_p = new_paragraphs[new_idx]
                last_seen_new_p = new_p

                # If text is 100% identical, leave paragraph completely untouched!
                # Preserves cover pages, images, drawings, formatting, line spacing, and section breaks.
                if old_p.text == new_p.text:
                    continue

                new_font_props = _extract_font_props(new_p)
                old_font_props = _extract_font_props(old_p)

                diff_chunks = self._compare_text(old_p.text, new_p.text)
                _clear_paragraph_runs(new_p)

                for chunk_tag, text in diff_chunks:
                    if not text:
                        continue
                    run = new_p.add_run(text)
                    if chunk_tag == 'equal':
                        _apply_font_props(run, new_font_props)
                    elif chunk_tag == 'insert':
                        _apply_font_props(run, new_font_props)
                        run.font.color.rgb = self.red_color
                        run.bold = True
                    elif chunk_tag == 'delete':
                        _apply_font_props(run, old_font_props)
                        run.font.color.rgb = self.red_color
                        run.font.strike = True

            elif tag == 'insert':
                new_p = new_paragraphs[new_idx]
                last_seen_new_p = new_p
                new_font_props = _extract_font_props(new_p)
                if not new_p.runs and new_p.text:
                    text = new_p.text
                    _clear_paragraph_runs(new_p)
                    run = new_p.add_run(text)
                    _apply_font_props(run, new_font_props)
                    run.font.color.rgb = self.red_color
                    run.bold = True
                else:
                    for run in new_p.runs:
                        run.font.color.rgb = self.red_color
                        run.bold = True

            elif tag == 'delete':
                if self.show_deletions:
                    old_p = old_paragraphs[old_idx]
                    if not old_p.text.strip():
                        continue

                    target_p = None
                    if new_idx is not None and new_idx < len(new_paragraphs):
                        target_p = new_paragraphs[new_idx]
                    elif last_seen_new_p is not None:
                        target_p = last_seen_new_p
                    elif len(new_paragraphs) > 0:
                        target_p = new_paragraphs[0]

                    if target_p is not None:
                        del_p = target_p.insert_paragraph_before(style=old_p.style)
                        old_font_props = _extract_font_props(old_p)
                        run = del_p.add_run(old_p.text)
                        _apply_font_props(run, old_font_props)
                        run.font.color.rgb = self.red_color
                        run.font.strike = True
                    else:
                        del_p = out_doc.add_paragraph(style=old_p.style)
                        old_font_props = _extract_font_props(old_p)
                        run = del_p.add_run(old_p.text)
                        _apply_font_props(run, old_font_props)
                        run.font.color.rgb = self.red_color
                        run.font.strike = True

        out_dir = os.path.dirname(os.path.abspath(self.output_path))
        if out_dir and not os.path.exists(out_dir):
            os.makedirs(out_dir, exist_ok=True)

        out_doc.save(self.output_path)
        return self.output_path

