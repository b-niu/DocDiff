import os
from docx import Document
from docdiff.core.engine import DocDiffEngine

def create_sample_docs():
    os.makedirs("tests/samples", exist_ok=True)
    old_path = "tests/samples/old_contract.docx"
    new_path = "tests/samples/new_contract.docx"
    out_path = "tests/samples/output_tracked.docx"

    # Old Document
    doc_old = Document()
    doc_old.add_heading("服务协议 (第一版)", level=1)
    doc_old.add_paragraph("甲方：张三科技有限公司")
    doc_old.add_paragraph("乙方：李四网络服务公司")
    doc_old.add_paragraph("1. 本协议服务费用为人民币 50,000 元整。")
    doc_old.add_paragraph("2. 乙方应当在 30 个工作日内完成交付。")
    doc_old.add_paragraph("3. 旧版特有条款：违约金为合同总金额的 5%。")
    doc_old.save(old_path)

    # New Document
    doc_new = Document()
    doc_new.add_heading("服务协议 (第二版更新)", level=1)
    doc_new.add_paragraph("甲方：张三科技有限公司")
    doc_new.add_paragraph("乙方：李四网络服务公司 (上海分公司)")
    doc_new.add_paragraph("1. 本协议服务费用为人民币 80,000 元整。")
    doc_new.add_paragraph("新增条款：服务期间提供 7x24 小时技术支持服务。")
    doc_new.add_paragraph("2. 乙方应当在 20 个工作日内完成交付。")
    doc_new.save(new_path)

    return old_path, new_path, out_path

def test_engine():
    old_path, new_path, out_path = create_sample_docs()
    engine = DocDiffEngine(old_path, new_path, out_path)
    res_path = engine.execute()
    print("Test finished successfully! Saved to:", res_path)
    assert os.path.exists(res_path)

    # Check output docx content
    doc_out = Document(res_path)
    print("\n--- Processed Document Paragraphs ---")
    for i, p in enumerate(doc_out.paragraphs):
        runs_info = [(r.text, r.font.color.rgb, r.font.strike) for r in p.runs]
        print(f"P[{i}]: {p.text} | Runs: {runs_info}")

if __name__ == "__main__":
    test_engine()
